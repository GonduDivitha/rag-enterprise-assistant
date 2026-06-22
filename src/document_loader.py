from pathlib import Path
from typing import List
import logging
import os

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader

from langchain_core.documents import Document

# ==========================================
# Logger Configuration
# ==========================================

logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    Loads PDF, DOCX, TXT, and CSV files
    and converts them into LangChain Documents.
    """

    @staticmethod
    def _get_metadata(file_path: str, file_type: str) -> dict:
        """
        Generate metadata for a document.
        """
        return {
            "source": Path(file_path).name,
            "file_type": file_type,
            "file_path": file_path,
            "file_size": os.path.getsize(file_path)
        }

    @staticmethod
    def load_pdf(file_path: str) -> List[Document]:
        """
        Load PDF and create page-level documents.
        """

        documents = []

        try:
            reader = PdfReader(file_path)

            for page_number, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text()

                if page_text and page_text.strip():

                    metadata = DocumentLoader._get_metadata(
                        file_path,
                        "pdf"
                    )

                    metadata["page"] = page_number

                    documents.append(
                        Document(
                            page_content=page_text,
                            metadata=metadata
                        )
                    )

            logger.info(
                f"Loaded PDF: {file_path} | Pages: {len(documents)}"
            )

            return documents

        except Exception as e:
            logger.error(
                f"Failed to load PDF {file_path}: {str(e)}"
            )
            raise

    @staticmethod
    def load_docx(file_path: str) -> List[Document]:
        """
        Load DOCX document.
        """

        try:
            doc = DocxDocument(file_path)

            text = "\n".join(
                paragraph.text
                for paragraph in doc.paragraphs
                if paragraph.text.strip()
            )

            if not text.strip():
                logger.warning(
                    f"Empty DOCX file: {file_path}"
                )

            document = Document(
                page_content=text,
                metadata=DocumentLoader._get_metadata(
                    file_path,
                    "docx"
                )
            )

            logger.info(
                f"Loaded DOCX: {file_path}"
            )

            return [document]

        except Exception as e:
            logger.error(
                f"Failed to load DOCX {file_path}: {str(e)}"
            )
            raise

    @staticmethod
    def load_txt(file_path: str) -> List[Document]:
        """
        Load TXT document with encoding fallback.
        """

        try:

            try:
                with open(
                    file_path,
                    "r",
                    encoding="utf-8"
                ) as file:
                    text = file.read()

            except UnicodeDecodeError:

                with open(
                    file_path,
                    "r",
                    encoding="latin-1"
                ) as file:
                    text = file.read()

            if not text.strip():
                logger.warning(
                    f"Empty TXT file: {file_path}"
                )

            document = Document(
                page_content=text,
                metadata=DocumentLoader._get_metadata(
                    file_path,
                    "txt"
                )
            )

            logger.info(
                f"Loaded TXT: {file_path}"
            )

            return [document]

        except Exception as e:
            logger.error(
                f"Failed to load TXT {file_path}: {str(e)}"
            )
            raise

    @staticmethod
    def load_csv(file_path: str) -> List[Document]:
        """
        Load CSV document.
        Large CSVs are limited to first 5000 rows
        to avoid memory issues.
        """

        try:
            dataframe = pd.read_csv(file_path)

            row_count = len(dataframe)

            if row_count > 5000:
                logger.warning(
                    f"Large CSV detected ({row_count} rows). "
                    f"Using first 5000 rows."
                )

                dataframe = dataframe.head(5000)

            text = dataframe.to_string(index=False)

            document = Document(
                page_content=text,
                metadata=DocumentLoader._get_metadata(
                    file_path,
                    "csv"
                )
            )

            logger.info(
                f"Loaded CSV: {file_path}"
            )

            return [document]

        except Exception as e:
            logger.error(
                f"Failed to load CSV {file_path}: {str(e)}"
            )
            raise

    @staticmethod
    def load_document(file_path: str) -> List[Document]:
        """
        Route file loading based on extension.
        """

        extension = Path(file_path).suffix.lower()

        if extension == ".pdf":
            return DocumentLoader.load_pdf(file_path)

        elif extension == ".docx":
            return DocumentLoader.load_docx(file_path)

        elif extension == ".txt":
            return DocumentLoader.load_txt(file_path)

        elif extension == ".csv":
            return DocumentLoader.load_csv(file_path)

        else:
            raise ValueError(
                f"Unsupported file type: {extension}"
            )